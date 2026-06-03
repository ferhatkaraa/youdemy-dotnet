from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "rapor"
IMG_DIR = OUT_DIR / "gorseller"
DOCX_PATH = OUT_DIR / "Youdemy_Proje_Raporu.docx"


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"


def ensure_dirs():
    IMG_DIR.mkdir(parents=True, exist_ok=True)


def load_font(size=26, bold=False):
    candidates = [
        "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
        "C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf",
    ]
    for path in candidates:
        if Path(path).exists():
            return ImageFont.truetype(path, size)
    return ImageFont.load_default()


def centered_text(draw, box, text, font, fill=(20, 30, 45)):
    x1, y1, x2, y2 = box
    bbox = draw.multiline_textbbox((0, 0), text, font=font, spacing=5, align="center")
    w = bbox[2] - bbox[0]
    h = bbox[3] - bbox[1]
    draw.multiline_text((x1 + (x2 - x1 - w) / 2, y1 + (y2 - y1 - h) / 2), text, font=font, fill=fill, spacing=5, align="center")


def arrow(draw, start, end, fill=(46, 116, 181), width=4):
    draw.line([start, end], fill=fill, width=width)
    sx, sy = start
    ex, ey = end
    if ex >= sx:
        head = [(ex, ey), (ex - 14, ey - 8), (ex - 14, ey + 8)]
    else:
        head = [(ex, ey), (ex + 14, ey - 8), (ex + 14, ey + 8)]
    draw.polygon(head, fill=fill)


def rounded_box(draw, box, text, fill, outline=(46, 116, 181), font_size=24):
    draw.rounded_rectangle(box, radius=18, fill=fill, outline=outline, width=3)
    centered_text(draw, box, text, load_font(font_size, bold=True))


def make_architecture_diagram():
    path = IMG_DIR / "mimari_sema.png"
    img = Image.new("RGB", (1400, 780), "white")
    d = ImageDraw.Draw(img)
    title = load_font(36, bold=True)
    d.text((60, 35), "Youdemy MVC Mimari Yapısı", font=title, fill=(20, 30, 45))

    rounded_box(d, (80, 140, 360, 260), "Kullanıcı\nTarayıcı", "#F8FAFC")
    rounded_box(d, (520, 120, 850, 240), "ASP.NET Core MVC\nController Katmanı", "#E8EEF5")
    rounded_box(d, (1010, 120, 1320, 240), "Razor View\nArayüz Katmanı", "#F8FAFC")
    rounded_box(d, (520, 340, 850, 460), "Model + ViewModel\nVeri Yapıları", "#F8FAFC")
    rounded_box(d, (1010, 340, 1320, 460), "ViewComponent\nKurs Müfredatı", "#F8FAFC")
    rounded_box(d, (520, 560, 850, 680), "EF Core\nDbContext", "#E8EEF5")
    rounded_box(d, (1010, 560, 1320, 680), "SQLite\nUygulama Verisi", "#F8FAFC")

    arrow(d, (360, 200), (520, 180))
    arrow(d, (850, 180), (1010, 180))
    arrow(d, (685, 240), (685, 340))
    arrow(d, (850, 400), (1010, 400))
    arrow(d, (685, 460), (685, 560))
    arrow(d, (850, 620), (1010, 620))
    arrow(d, (1010, 200), (360, 220), fill=(100, 116, 139))

    d.text((85, 715), "Şekil 1. Uygulama, video dosyalarını değil; kullanıcı, kurs, ders, abonelik ve ilerleme verilerini yönetir.", font=load_font(20), fill=(70, 80, 95))
    img.save(path)
    return path


def make_flow_diagram():
    path = IMG_DIR / "is_akisi.png"
    img = Image.new("RGB", (1400, 760), "white")
    d = ImageDraw.Draw(img)
    d.text((60, 35), "Temel Kullanım Akışı", font=load_font(36, bold=True), fill=(20, 30, 45))

    boxes = [
        ((80, 160, 330, 270), "Eğitmen\nkurs oluşturur"),
        ((410, 160, 660, 270), "Ders linki\nekler"),
        ((740, 160, 990, 270), "Öğrenci\nkursa abone olur"),
        ((1070, 160, 1320, 270), "Dersi izler\nve işaretler"),
    ]
    for box, text in boxes:
        rounded_box(d, box, text, "#F8FAFC", font_size=23)
    for i in range(len(boxes) - 1):
        arrow(d, (boxes[i][0][2], 215), (boxes[i + 1][0][0], 215))

    rounded_box(d, (210, 430, 530, 560), "CourseAuthorize\nabonelik kontrolü", "#E8EEF5", font_size=23)
    rounded_box(d, (630, 430, 950, 560), "CourseSyllabus\nViewComponent", "#E8EEF5", font_size=23)
    rounded_box(d, (1050, 430, 1300, 560), "LessonProgress\ncheckbox durumu", "#E8EEF5", font_size=23)
    arrow(d, (865, 270), (370, 430), fill=(100, 116, 139))
    arrow(d, (1195, 270), (790, 430), fill=(100, 116, 139))
    arrow(d, (1195, 270), (1175, 430), fill=(100, 116, 139))

    d.text((85, 690), "Şekil 2. Öğrenci kursu izleyebilmek için önce abone olur; ilerleme durumu ders bazında saklanır.", font=load_font(20), fill=(70, 80, 95))
    img.save(path)
    return path


def make_er_diagram():
    path = IMG_DIR / "veritabani_sema.png"
    img = Image.new("RGB", (1400, 820), "white")
    d = ImageDraw.Draw(img)
    d.text((60, 35), "Veritabanı İlişkileri", font=load_font(36, bold=True), fill=(20, 30, 45))

    boxes = [
        ((80, 150, 360, 300), "User\nId, Username\nRole, IsApproved"),
        ((560, 150, 860, 300), "Course\nId, Title\nTeacherId"),
        ((1040, 150, 1320, 300), "Lesson\nId, Title\nVideoUrl, CourseId"),
        ((320, 480, 640, 630), "Enrollment\nStudentId\nCourseId"),
        ((820, 480, 1160, 630), "LessonProgress\nStudentId\nLessonId\nIsCompleted"),
    ]
    for box, text in boxes:
        rounded_box(d, box, text, "#F8FAFC", font_size=22)

    arrow(d, (360, 215), (560, 215))
    d.text((405, 185), "1 eğitmen\nN kurs", font=load_font(18), fill=(70, 80, 95))
    arrow(d, (860, 215), (1040, 215))
    d.text((910, 185), "1 kurs\nN ders", font=load_font(18), fill=(70, 80, 95))
    arrow(d, (240, 300), (420, 480), fill=(100, 116, 139))
    arrow(d, (710, 300), (500, 480), fill=(100, 116, 139))
    arrow(d, (240, 300), (990, 480), fill=(100, 116, 139))
    arrow(d, (1180, 300), (990, 480), fill=(100, 116, 139))

    d.text((85, 735), "Şekil 3. Veritabanı video dosyasını değil; link, kayıt ve ilerleme bilgisini tutar.", font=load_font(20), fill=(70, 80, 95))
    img.save(path)
    return path


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Sayfa ")
    run.font.size = Pt(9)
    run.font.name = "Calibri"

    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)

    run = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    run._r.append(instr)

    run = paragraph.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_end)


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.49)
    section.footer_distance = Inches(0.49)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    footer = section.footer.paragraphs[0]
    add_page_number(footer)


def add_title_page(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(120)
    r = p.add_run("Youdemy (.NET) v2\n")
    r.bold = True
    r.font.size = Pt(26)
    r.font.color.rgb = BLUE
    r.font.name = "Calibri"
    r2 = p.add_run("Online Kurs / LMS Platformu Proje Raporu")
    r2.bold = True
    r2.font.size = Pt(20)
    r2.font.name = "Calibri"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(30)
    p.add_run("Mini-Udemy benzeri kurs yönetimi, ders linki ekleme, abonelik ve ilerleme takip sistemi").italic = True

    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    rows = [
        ("Proje Türü", "ASP.NET Core MVC web uygulaması"),
        ("Veritabanı", "SQLite + Entity Framework Core"),
        ("Temel Yaklaşım", "Video dosyası saklamak yerine dış servis linklerini kullanma"),
        ("Rapor Kapsamı", "Analiz, mimari, veritabanı, yetkilendirme, arayüz ve test değerlendirmesi"),
    ]
    for row, (k, v) in zip(table.rows, rows):
        set_cell_text(row.cells[0], k, bold=True)
        set_cell_text(row.cells[1], v)
        set_cell_shading(row.cells[0], LIGHT_BLUE)

    doc.add_paragraph()
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.add_run("Hazırlanan bu rapor, proje teslim dokümanı formatında düzenlenmiştir.").bold = True
    doc.add_page_break()


def add_toc(doc):
    doc.add_heading("İçindekiler", level=1)
    items = [
        ("1. Giriş", ["1.1. Projenin Tanımı", "1.2. Projenin Amacı", "1.3. Projenin Kapsamı"]),
        ("2. Kullanılan Teknolojiler", ["2.1. ASP.NET Core MVC", "2.2. Entity Framework Core ve SQLite", "2.3. HTML, CSS ve JavaScript"]),
        ("3. Sistem Analizi", ["3.1. Kullanıcı Rolleri", "3.2. Eğitmen Gereksinimleri", "3.3. Öğrenci Gereksinimleri", "3.4. Yönetici Gereksinimleri"]),
        ("4. Veritabanı Tasarımı", ["4.1. Temel Tablolar", "4.2. İlişkiler", "4.3. Video Linki Saklama Yaklaşımı"]),
        ("5. Proje Mimarisi", ["5.1. MVC Yapısı", "5.2. Controller ve Model Katmanı", "5.3. ViewComponent ve Helper Yapıları"]),
        ("6. Kurs ve Ders Yönetimi", ["6.1. Kurs Oluşturma", "6.2. Ders Ekleme", "6.3. Video/Ders Linki Normalizasyonu"]),
        ("7. Öğrenci İşlemleri", ["7.1. Kursa Abone Olma", "7.2. Ders İzleme", "7.3. Checkbox ile İlerleme Takibi"]),
        ("8. Yetkilendirme ve Erişim Kontrolü", ["8.1. Rol Tabanlı Yetkilendirme", "8.2. CourseAuthorizeAttribute", "8.3. Güvenlik Değerlendirmesi"]),
        ("9. Arayüz ve Bileşen Kullanımı", ["9.1. Sayfa Yapıları", "9.2. CourseSyllabus ViewComponent", "9.3. Görsel Yerleşim Notları"]),
        ("10. Test, Değerlendirme ve Sonuç", ["10.1. Derleme Kontrolleri", "10.2. Test Süreci", "10.3. Genel Değerlendirme", "10.4. Sonuç"]),
    ]
    for title, subs in items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.add_run(title).bold = True
        for sub in subs:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.7)
            p.paragraph_format.space_after = Pt(1)
            p.add_run(sub)
    doc.add_page_break()


def add_placeholder(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, "FFF7E6")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"({text})")
    run.bold = True
    run.italic = True
    run.font.size = Pt(10)
    doc.add_paragraph()


def add_image(doc, path, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(6.2))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].italic = True
    cap.runs[0].font.size = Pt(9)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, bold=True)
        set_cell_shading(table.rows[0].cells[i], LIGHT_GRAY)
    for row_data in rows:
        row = table.add_row()
        for i, value in enumerate(row_data):
            set_cell_text(row.cells[i], value)
    doc.add_paragraph()


def body_paragraph(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run(text)
    return p


def build_report():
    ensure_dirs()
    architecture = make_architecture_diagram()
    flow = make_flow_diagram()
    er = make_er_diagram()

    doc = Document()
    configure_document(doc)
    add_title_page(doc)
    add_toc(doc)

    doc.add_heading("1. Giriş", level=1)
    doc.add_heading("1.1. Projenin Tanımı", level=2)
    body_paragraph(doc, "Youdemy (.NET) v2 projesi, çevrim içi eğitim içeriklerinin düzenlenmesi ve öğrencilere sunulması amacıyla geliştirilmiş bir Online Kurs / LMS platformudur. Proje, Udemy benzeri temel işlevleri sade ve anlaşılır bir yapı içinde sunmayı hedefler. Eğitmenler kurs oluşturabilir, ders başlığı ve açıklaması ekleyebilir, ders videosu için dış kaynak linki tanımlayabilir. Öğrenciler ise ilgilendikleri kurslara abone olarak dersleri izleyebilir ve izleme durumlarını checkbox aracılığıyla takip edebilir.")
    body_paragraph(doc, "Uygulamanın en önemli tasarım kararlarından biri video dosyalarını uygulama sunucusunda veya veritabanında saklamamasıdır. Bunun yerine YouTube, Google Drive ve Vimeo gibi servislerde tutulan videolara ait linkler sisteme eklenir. Böylece büyük dosya depolama, bant genişliği ve medya sunucusu yönetimi gibi maliyetler azaltılır.")
    doc.add_heading("1.2. Projenin Amacı", level=2)
    body_paragraph(doc, "Projenin amacı, eğitmen ve öğrenci rollerini içeren temel bir LMS sistemi oluşturmaktır. Bu sistemde eğitmen tarafında kurs ve ders yönetimi, öğrenci tarafında abonelik ve ilerleme takibi, sistem tarafında ise güvenli erişim kontrolü sağlanır. Proje aynı zamanda ASP.NET Core MVC, Entity Framework Core, ViewComponent kullanımı ve Authorize yapılarının pratik bir örnek üzerinde uygulanmasını sağlar.")
    doc.add_heading("1.3. Projenin Kapsamı", level=2)
    add_bullets(doc, [
        "Eğitmenlerin kurs oluşturması ve ders linkleri eklemesi.",
        "Öğrencilerin kurslara abone olması ve yalnızca abone oldukları kursları izlemesi.",
        "Ders ilerlemesinin checkbox ile işaretlenmesi ve veritabanında saklanması.",
        "Kurs müfredatının ViewComponent yapısı ile tekrar kullanılabilir şekilde gösterilmesi.",
        "Rol tabanlı yetkilendirme ve özel kurs erişim filtresi kullanılması.",
    ])
    add_placeholder(doc, "Buraya projenin ana sayfasına ait bir ekran görüntüsü gelecek. Görselde kurs kartları ve genel LMS arayüzü görünmelidir.")
    doc.add_page_break()

    doc.add_heading("2. Kullanılan Teknolojiler", level=1)
    doc.add_heading("2.1. ASP.NET Core MVC", level=2)
    body_paragraph(doc, "Proje ASP.NET Core MVC mimarisiyle geliştirilmiştir. MVC yapısı, uygulamanın veri modeli, iş akışı ve arayüzünü birbirinden ayırarak daha düzenli bir proje yapısı sağlar. Controller dosyaları kullanıcı isteklerini karşılar, modeller veri yapısını temsil eder, view dosyaları ise Razor sözdizimiyle HTML çıktısı üretir.")
    doc.add_heading("2.2. Entity Framework Core ve SQLite", level=2)
    body_paragraph(doc, "Veri yönetimi için Entity Framework Core kullanılmıştır. EF Core, C# sınıfları ile veritabanı tabloları arasında ilişki kurmayı sağlar. SQLite ise hafif, dosya tabanlı ve eğitim projeleri için uygun bir veritabanı çözümüdür. Bu projede kullanıcılar, kurslar, dersler, abonelikler, bildirimler ve ders ilerleme kayıtları SQLite üzerinde tutulur.")
    doc.add_heading("2.3. HTML, CSS ve JavaScript", level=2)
    body_paragraph(doc, "Arayüz tarafında Razor view yapıları HTML ile birlikte kullanılmıştır. CSS dosyaları sayfaların görsel düzenini ve kart yapısını belirler. JavaScript ise özellikle ders ilerleme checkbox işlemlerinde AJAX isteği göndermek ve sayfadaki ilerleme çubuğunu anlık güncellemek için kullanılır.")
    add_table(doc, ["Teknoloji", "Projede Kullanım Amacı"], [
        ("ASP.NET Core MVC", "Web uygulamasının ana mimari yapısı"),
        ("Entity Framework Core", "Veritabanı işlemlerinin nesne yönelimli yönetimi"),
        ("SQLite", "Kullanıcı, kurs, ders ve ilerleme verilerinin saklanması"),
        ("Razor Views", "Dinamik HTML sayfalarının oluşturulması"),
        ("JavaScript", "Checkbox ilerleme güncellemesi ve küçük etkileşimler"),
    ])
    doc.add_page_break()

    doc.add_heading("3. Sistem Analizi", level=1)
    doc.add_heading("3.1. Kullanıcı Rolleri", level=2)
    body_paragraph(doc, "Sistemde üç temel kullanıcı rolü bulunmaktadır: Student, Teacher ve Admin. Student rolü öğrencileri temsil eder ve kurslara abone olma, ders izleme ve ilerleme işaretleme işlemlerini gerçekleştirir. Teacher rolü eğitmenleri temsil eder ve kurs oluşturma, ders ekleme ve kendi kursunu önizleme yetkilerine sahiptir. Admin rolü ise sistem yönetimi ve kullanıcı onay süreçlerinde kullanılır.")
    doc.add_heading("3.2. Eğitmen Gereksinimleri", level=2)
    add_bullets(doc, [
        "Kurs başlığı, açıklaması ve kapak görseli ekleyebilmelidir.",
        "Kursa bağlı dersler oluşturabilmelidir.",
        "Derslere video veya ders linki ekleyebilmelidir.",
        "Yalnızca kendi kurslarını düzenleyebilmelidir.",
        "Hesabı onaylanmadan kurs ekleme işlemi yapamamalıdır.",
    ])
    doc.add_heading("3.3. Öğrenci Gereksinimleri", level=2)
    add_bullets(doc, [
        "Katalogdaki kursları inceleyebilmelidir.",
        "Kursa abone olduktan sonra izleme sayfasına erişebilmelidir.",
        "Dersleri sırayla görüntüleyebilmelidir.",
        "Tamamladığı dersleri checkbox ile işaretleyebilmelidir.",
        "Kendi kurs ilerleme yüzdesini görebilmelidir.",
    ])
    doc.add_heading("3.4. Yönetici Gereksinimleri", level=2)
    body_paragraph(doc, "Yönetici rolü, öğretmen hesaplarının onaylanması ve sistem genelindeki yönetim işlemleri için kullanılır. Bu yapı, eğitmenlerin doğrudan içerik yayınlamadan önce kontrol edilebilmesine olanak sağlar. Eğitim platformlarında içerik güvenliği ve kalite kontrol açısından bu tür bir onay mekanizması önemlidir.")
    add_image(doc, flow, "Şekil 1. Temel kullanım akışı")
    doc.add_page_break()

    doc.add_heading("4. Veritabanı Tasarımı", level=1)
    doc.add_heading("4.1. Temel Tablolar", level=2)
    body_paragraph(doc, "Veritabanı tasarımı, platformun ihtiyaç duyduğu uygulama verilerini saklamak için oluşturulmuştur. Projede video dosyaları saklanmaz; yalnızca video/ders linkleri ders tablosunda tutulur. Bu yaklaşım, veritabanının boyutunu küçük tutar ve medya depolama yükünü dış servislere devreder.")
    add_table(doc, ["Tablo", "Görev"], [
        ("Users", "Öğrenci, eğitmen ve admin kullanıcı bilgilerini tutar."),
        ("Courses", "Kurs başlığı, açıklaması, görseli ve eğitmen ilişkisini tutar."),
        ("Lessons", "Ders başlığı, açıklaması, sırası ve video linkini tutar."),
        ("Enrollments", "Öğrencinin hangi kursa abone olduğunu tutar."),
        ("LessonProgresses", "Öğrencinin ders tamamlama durumunu tutar."),
        ("Notifications", "Sistem veya admin tarafından gönderilen bildirimleri tutar."),
    ])
    doc.add_heading("4.2. İlişkiler", level=2)
    body_paragraph(doc, "Course ile User arasında eğitmen ilişkisi bulunur. Bir eğitmen birden fazla kurs oluşturabilir. Course ile Lesson arasında bire-çok ilişki vardır; bir kurs birden fazla dersten oluşur. Enrollment tablosu öğrenci ile kurs arasındaki abonelik ilişkisini kurar. LessonProgress tablosu ise öğrenci ile ders arasındaki tamamlanma bilgisini saklar.")
    add_image(doc, er, "Şekil 2. Veritabanı ilişki şeması")
    doc.add_heading("4.3. Video Linki Saklama Yaklaşımı", level=2)
    body_paragraph(doc, "Video dosyaları büyük boyutlu olduğu için doğrudan veritabanında tutulmaları uygun değildir. Bu projede ders modeli içinde VideoUrl alanı yer alır. Eğitmen YouTube, Google Drive veya Vimeo gibi bir kaynakta bulunan videonun linkini girer. Uygulama bu linki izleme ekranına uygun embed veya preview formatına dönüştürür.")
    doc.add_page_break()

    doc.add_heading("5. Proje Mimarisi", level=1)
    doc.add_heading("5.1. MVC Yapısı", level=2)
    body_paragraph(doc, "MVC mimarisi, projenin bakımını ve anlaşılmasını kolaylaştırır. Controller katmanı istekleri alır ve gerekli verileri model katmanından çeker. View katmanı kullanıcıya gösterilecek HTML çıktısını üretir. Model katmanı ise uygulama verisinin şeklini ve ilişkilerini tanımlar.")
    doc.add_heading("5.2. Controller ve Model Katmanı", level=2)
    body_paragraph(doc, "CoursesController kurs listeleme, detay görüntüleme, abonelik, izleme ve yönetim işlemlerini içerir. LessonsController ders oluşturma, düzenleme, silme ve ilerleme durumunu değiştirme işlemlerini yürütür. Modeller tarafında Course, Lesson, Enrollment ve LessonProgress sınıfları LMS sisteminin temelini oluşturur.")
    doc.add_heading("5.3. ViewComponent ve Helper Yapıları", level=2)
    body_paragraph(doc, "CourseSyllabusViewComponent, izleme ekranındaki kurs müfredatını ve tamamlanan ders durumlarını göstermek için kullanılır. VideoLinkHelper ise eğitmen tarafından girilen farklı video linklerini normalize ederek oynatılabilir hale getirir. Bu yardımcı yapı, link dönüştürme işini controller içinden ayırarak kodun daha temiz olmasını sağlar.")
    add_image(doc, architecture, "Şekil 3. ASP.NET Core MVC mimari yapısı")
    doc.add_page_break()

    doc.add_heading("6. Kurs ve Ders Yönetimi", level=1)
    doc.add_heading("6.1. Kurs Oluşturma", level=2)
    body_paragraph(doc, "Eğitmen hesabı onaylandıktan sonra kurs oluşturabilir. Kurs oluşturma ekranında kurs başlığı, açıklaması ve isteğe bağlı kapak görseli bilgisi girilir. Kurs kaydedildikten sonra öğretmen kendi yönetim ekranı üzerinden kurslarını görebilir ve düzenleyebilir.")
    doc.add_heading("6.2. Ders Ekleme", level=2)
    body_paragraph(doc, "Her kurs birden fazla dersten oluşabilir. Ders ekleme ekranında ders başlığı, açıklaması, ders sırası ve video/ders linki alanları bulunur. Ders sırası, müfredatın izleme ekranında hangi sırayla gösterileceğini belirler. Bu yapı, öğrencinin kursu daha düzenli takip etmesini sağlar.")
    doc.add_heading("6.3. Video/Ders Linki Normalizasyonu", level=2)
    body_paragraph(doc, "VideoLinkHelper sınıfı, eğitmen tarafından girilen YouTube, youtu.be, YouTube Shorts, Google Drive, Vimeo veya iframe embed kodlarını tanır. Link uygun formatta değilse izleme ekranında kullanılabilecek embed/preview formatına dönüştürülür. Böylece eğitmen farklı kaynaklardan video ekleyebilir.")
    add_table(doc, ["Girilen Link Türü", "Sistemin Davranışı"], [
        ("YouTube watch linki", "Embed formatına dönüştürülür."),
        ("youtu.be kısa linki", "YouTube embed formatına dönüştürülür."),
        ("YouTube Shorts", "Standart embed formatına dönüştürülür."),
        ("Google Drive dosya linki", "Preview formatına dönüştürülür."),
        ("Vimeo linki", "Player linkine dönüştürülür."),
        ("iframe kodu", "src içindeki link alınır."),
    ])
    add_placeholder(doc, "Buraya eğitmen ders ekleme ekranının ekran görüntüsü gelecek. Görselde Video / Ders Linki alanı görünmelidir.")
    doc.add_page_break()

    doc.add_heading("7. Öğrenci İşlemleri", level=1)
    doc.add_heading("7.1. Kursa Abone Olma", level=2)
    body_paragraph(doc, "Öğrenciler kurs detay sayfasından kursa abone olabilir. Abonelik işlemi Enrollment tablosuna kayıt eklenmesiyle gerçekleşir. Öğrenci daha önce aynı kursa abone olmuşsa tekrar kayıt oluşturulmaz. Bu durum veritabanında öğrenci ve kurs için benzersiz indeks ile de desteklenir.")
    doc.add_heading("7.2. Ders İzleme", level=2)
    body_paragraph(doc, "Abone olan öğrenci kurs izleme sayfasına erişebilir. İzleme sayfasında sol tarafta aktif ders videosu ve ders açıklaması, sağ tarafta ise kurs müfredatı ve ilerleme bilgisi bulunur. Öğrenci müfredat listesindeki derslere tıklayarak aktif dersi değiştirebilir.")
    doc.add_heading("7.3. Checkbox ile İlerleme Takibi", level=2)
    body_paragraph(doc, "Her dersin yanında yer alan checkbox öğrencinin o dersi tamamlayıp tamamlamadığını temsil eder. Checkbox tıklandığında JavaScript ile Lessons/ToggleComplete endpoint'ine istek gönderilir. Sunucu tarafında öğrencinin kursa kayıtlı olup olmadığı kontrol edilir ve LessonProgress kaydı güncellenir.")
    body_paragraph(doc, "İlerleme bilgisi güncellendikten sonra tamamlanan ders sayısı, toplam ders sayısı ve yüzde değeri hesaplanarak JSON olarak geri döndürülür. Sayfa yenilenmeden ilerleme çubuğu ve sayaçlar güncellenir. Bu yapı kullanıcı deneyimini daha akıcı hale getirir.")
    add_placeholder(doc, "Buraya öğrenci ders izleme ekranının ekran görüntüsü gelecek. Görselde video alanı, müfredat listesi, checkbox ve ilerleme çubuğu görünmelidir.")
    doc.add_page_break()

    doc.add_heading("8. Yetkilendirme ve Erişim Kontrolü", level=1)
    doc.add_heading("8.1. Rol Tabanlı Yetkilendirme", level=2)
    body_paragraph(doc, "Projedeki işlemler kullanıcı rollerine göre sınırlandırılmıştır. Örneğin kurs oluşturma ve ders ekleme işlemleri yalnızca Teacher rolüne sahip kullanıcılar tarafından yapılabilir. Öğrenciye ait kurslarım ve ilerleme işlemleri Student rolü ile sınırlandırılmıştır. Bu ayrım, sistemde yetkisiz işlem yapılmasını engeller.")
    doc.add_heading("8.2. CourseAuthorizeAttribute", level=2)
    body_paragraph(doc, "Kurs izleme sayfasına erişim için özel CourseAuthorizeAttribute filtresi kullanılmıştır. Bu filtre kullanıcının giriş yapıp yapmadığını, rolünü ve ilgili kursa erişim hakkını kontrol eder. Öğrenciler yalnızca abone oldukları kursu izleyebilir. Eğitmen yalnızca kendi kursunu önizleyebilir. Admin rolü ise sistem yönetimi amacıyla erişime sahiptir.")
    doc.add_heading("8.3. Güvenlik Değerlendirmesi", level=2)
    body_paragraph(doc, "Kullanıcı kimlik doğrulama işlemleri Cookie Authentication ile yapılır. Kritik POST işlemlerinde anti-forgery token kullanılması CSRF saldırılarına karşı koruma sağlar. Ders ilerleme güncellemesi de ValidateAntiForgeryToken ile korunmuştur. Ayrıca veritabanında tekrar eden abonelik ve ilerleme kayıtlarını önlemek için benzersiz indeksler tanımlanmıştır.")
    add_table(doc, ["Kontrol", "Uygulamadaki Karşılığı"], [
        ("Rol kontrolü", "[Authorize(Roles = \"Teacher\")] ve [Authorize(Roles = \"Student\")] kullanımı"),
        ("Kurs erişimi", "CourseAuthorizeAttribute ile abonelik ve sahiplik kontrolü"),
        ("CSRF koruması", "POST işlemlerinde ValidateAntiForgeryToken kullanımı"),
        ("Tekrarlı kayıt önleme", "Enrollment ve LessonProgress üzerinde benzersiz indeksler"),
    ])
    doc.add_page_break()

    doc.add_heading("9. Arayüz ve Bileşen Kullanımı", level=1)
    doc.add_heading("9.1. Sayfa Yapıları", level=2)
    body_paragraph(doc, "Uygulamada ana sayfa, kurs detay sayfası, öğretmen kurs yönetim ekranı, ders ekleme/düzenleme ekranları, öğrenci kurslarım ekranı ve ders izleme ekranı gibi temel sayfalar bulunur. Arayüz, kurs kartları ve içerik panelleri ile kullanıcıların sistemi rahat kullanmasını hedefler.")
    doc.add_heading("9.2. CourseSyllabus ViewComponent", level=2)
    body_paragraph(doc, "CourseSyllabus ViewComponent, ders izleme ekranının sağ tarafında yer alan müfredat listesini üretir. Aktif ders, tamamlanan dersler ve öğrenci rolü bilgisine göre farklı görsel durumlar oluşturulur. Bu bileşenin ayrı bir sınıfta bulunması, aynı müfredat gösteriminin farklı sayfalarda da kullanılabilmesine imkan verir.")
    doc.add_heading("9.3. Görsel Yerleşim Notları", level=2)
    body_paragraph(doc, "Rapor içerisinde mimari şema, iş akışı ve veritabanı ilişki şeması oluşturulmuştur. Ekran görüntüsü gerektiren alanlarda ise doğrudan uygulamadan alınması gereken görseller için yer tutucular bırakılmıştır. Bu sayede rapor tesliminden önce gerçek uygulama ekran görüntüleri kolayca eklenebilir.")
    add_placeholder(doc, "Buraya öğretmen kurs yönetim ekranının ekran görüntüsü gelecek. Görselde kurs listesi, ders ekleme bağlantısı ve yönetim butonları görünmelidir.")
    add_placeholder(doc, "Buraya öğrenci Kurslarım sayfasının ekran görüntüsü gelecek. Görselde abone olunan kurslar ve ilerleme yüzdesi görünmelidir.")
    doc.add_page_break()

    doc.add_heading("10. Test, Değerlendirme ve Sonuç", level=1)
    doc.add_heading("10.1. Derleme Kontrolleri", level=2)
    body_paragraph(doc, "Projede yapılan geliştirmelerden sonra dotnet build komutu çalıştırılmıştır. Son kontrolde proje 0 uyarı ve 0 hata ile başarılı şekilde derlenmiştir. Bu sonuç, C# kodu, Razor view dosyaları ve proje bağımlılıklarının derleme seviyesinde tutarlı olduğunu gösterir.")
    doc.add_heading("10.2. Test Süreci", level=2)
    body_paragraph(doc, "dotnet test komutu çalıştırılmış ve komut hatasız tamamlanmıştır. Projede ayrı bir test projesi bulunmadığı için bu komut mevcut proje yapılandırmasının test edilebilirlik kontrolünü yapmıştır. Fonksiyonel doğrulama için kullanıcı rolleriyle giriş yapılarak kurs oluşturma, ders ekleme, abonelik, izleme ve checkbox ile ilerleme senaryolarının elle denenmesi önerilir.")
    add_table(doc, ["Test Senaryosu", "Beklenen Sonuç"], [
        ("Eğitmen kurs oluşturur", "Kurs yönetim ekranında yeni kurs görünür."),
        ("Eğitmen ders linki ekler", "Ders kurs detayında ve izleme ekranında görünür."),
        ("Öğrenci kursa abone olur", "Öğrenci izleme sayfasına erişebilir."),
        ("Abone olmayan öğrenci izlemeye çalışır", "Kurs detay sayfasına yönlendirilir ve uyarı gösterilir."),
        ("Checkbox işaretlenir", "Ders tamamlandı görünür ve ilerleme yüzdesi güncellenir."),
    ])
    doc.add_heading("10.3. Genel Değerlendirme", level=2)
    body_paragraph(doc, "Youdemy projesi, verilen Online Kurs / LMS Platformu konusunun temel gereksinimlerini karşılamaktadır. Eğitmenlerin kurs ve ders linki ekleyebilmesi, öğrencilerin kurslara abone olup ilerleme takibi yapabilmesi, ViewComponent kullanımı ve Authorize attribute'ları ile erişim kontrolü sağlanması projenin ana hedefleriyle uyumludur.")
    body_paragraph(doc, "Projenin güçlü yönlerinden biri video saklama sorununun dış servis linkleriyle çözülmesidir. Bu yaklaşım daha az depolama ihtiyacı oluşturur ve gerçek hayattaki eğitim platformlarında sık kullanılan bir modele yakındır. Ayrıca Google Drive ve Vimeo gibi farklı kaynakların desteklenmesi sistemi daha esnek hale getirir.")
    doc.add_heading("10.4. Sonuç", level=2)
    body_paragraph(doc, "Sonuç olarak Youdemy (.NET) v2, ASP.NET Core MVC ile geliştirilen, rol tabanlı yetkilendirme kullanan ve temel LMS özelliklerini içeren başarılı bir web uygulamasıdır. Proje, hem yazılım mimarisi hem de kullanıcı senaryoları açısından eğitim amaçlı bir Mini-Udemy örneği olarak değerlendirilebilir. Geliştirme sürecinde veritabanı, controller, view, ViewComponent, helper ve filtre yapıları birlikte kullanılarak bütünlüklü bir sistem ortaya çıkarılmıştır.")
    body_paragraph(doc, "İlerleyen aşamalarda sisteme yorum ve puanlama, sertifika oluşturma, gelişmiş arama/filtreleme, kategori yönetimi ve daha kapsamlı test projesi eklenebilir. Ancak mevcut haliyle proje, ödevde belirtilen ana gereksinimleri karşılayacak düzeydedir.")

    doc.save(DOCX_PATH)
    return DOCX_PATH


if __name__ == "__main__":
    path = build_report()
    print(path)
